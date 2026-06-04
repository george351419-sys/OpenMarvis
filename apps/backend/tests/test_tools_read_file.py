"""read_file 工具：覆盖各格式 + 分页 + 大小上限 + 拒绝未知扩展名。"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from openmarvis.llm.event_sink import QueueEventSink
from openmarvis.security.policy import SecurityGate
from openmarvis.store.db import create_engine, init_db
from openmarvis.tools.base import ToolContext
from openmarvis.tools.read_file import ReadFileTool
from openmarvis.workspace.manager import Workspace


@pytest.fixture
def ctx(tmp_path):
    ws = Workspace(conv_id="c", root_base=tmp_path)
    ws.ensure()
    engine = create_engine(tmp_path / "db.sqlite")
    init_db(engine)
    return ToolContext(
        conv_id="c", agent_id="main", workspace=ws,
        memory_store=None,
        security=SecurityGate(workspace=ws),
        event_sink=QueueEventSink(),
        user_settings=None,
    )


async def test_read_plain_txt(ctx):
    f = ctx.workspace.output_dir / "a.txt"
    f.write_text("line1\nline2\nline3", encoding="utf-8")
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f)), ctx,
    )
    assert r.error is None
    assert "line1" in r.content
    assert "line3" in r.content


async def test_read_json_pretty_prints(ctx):
    f = ctx.workspace.output_dir / "a.json"
    f.write_text('{"b":1,"a":2}', encoding="utf-8")
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f)), ctx,
    )
    assert r.error is None
    # 美化后应有缩进
    assert "  " in r.content


async def test_read_csv_to_markdown_table(ctx):
    f = ctx.workspace.output_dir / "data.csv"
    f.write_text("name,age\nAlice,30\nBob,25\n", encoding="utf-8")
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f)), ctx,
    )
    assert r.error is None
    assert "| name | age |" in r.content
    assert "| Alice | 30 |" in r.content


async def test_read_docx(ctx):
    from docx import Document
    f = ctx.workspace.output_dir / "doc.docx"
    d = Document()
    d.add_heading("Title", level=1)
    d.add_paragraph("Hello world.")
    d.save(str(f))
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f)), ctx,
    )
    assert r.error is None
    assert "# Title" in r.content
    assert "Hello world" in r.content


async def test_read_xlsx_default_first_sheet(ctx):
    from openpyxl import Workbook
    f = ctx.workspace.output_dir / "x.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["a", "b"])
    ws.append([1, 2])
    wb.create_sheet("Other").append(["x"])
    wb.save(str(f))
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f)), ctx,
    )
    assert r.error is None
    assert "Sheet1" in r.content
    assert "Other" not in r.content


async def test_read_xlsx_read_all_sheets(ctx):
    from openpyxl import Workbook
    f = ctx.workspace.output_dir / "x.xlsx"
    wb = Workbook()
    wb.active.title = "S1"
    wb.active.append(["a"])
    wb.create_sheet("S2").append(["b"])
    wb.save(str(f))
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f), read_all_sheets=True), ctx,
    )
    assert r.error is None
    assert "S1" in r.content and "S2" in r.content


async def test_rejects_unknown_extension(ctx):
    f = ctx.workspace.output_dir / "blob.xyz"
    f.write_bytes(b"\x00\x01")
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f)), ctx,
    )
    assert r.error is not None
    assert ".xyz" in r.error


async def test_rejects_too_large(ctx):
    # 直接 patch MAX_BYTES 太麻烦——用真实大文件验证。
    # 写 ~1MB 但临时下调 limit；这里用 patch
    import openmarvis.tools.read_file as rf_mod
    f = ctx.workspace.output_dir / "big.txt"
    f.write_bytes(b"x" * 2000)  # 2KB
    old = rf_mod.MAX_BYTES
    rf_mod.MAX_BYTES = 1000  # 临时调到 1KB
    try:
        r = await ReadFileTool().execute(
            ReadFileTool.args_model(file_path=str(f)), ctx,
        )
        assert r.error is not None
        assert "过大" in r.error
    finally:
        rf_mod.MAX_BYTES = old


async def test_pagination(ctx):
    f = ctx.workspace.output_dir / "many.txt"
    f.write_text("\n".join(f"line{i}" for i in range(100)), encoding="utf-8")
    r = await ReadFileTool().execute(
        ReadFileTool.args_model(file_path=str(f), offset=10, limit=5),
        ctx,
    )
    assert r.error is None
    assert "line10" in r.content
    assert "line14" in r.content
    assert "line15" not in r.content.split("...")[0]  # 截断符前不应有 line15
    assert "已截断" in r.content


async def test_alias_path_param(ctx):
    # Hunyuan 之类的模型有时用 'path' 而非 'file_path'
    f = ctx.workspace.output_dir / "a.txt"
    f.write_text("hi", encoding="utf-8")
    args = ReadFileTool.args_model.model_validate({"path": str(f)})
    assert args.file_path == str(f)
    r = await ReadFileTool().execute(args, ctx)
    assert r.error is None
    assert "hi" in r.content
